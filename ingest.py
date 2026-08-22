"""
Đông Đô CS Chatbot - Document Ingestion Script
Đọc tất cả file .docx từ thư mục tailieu → chunking → embedding → ChromaDB
"""
import os
import glob
from datetime import datetime

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_chroma import Chroma

from config import (
    DOCUMENTS_DIR,
    VECTORDB_DIR,
    EMBEDDING_MODEL,
    CHROMA_COLLECTION_NAME,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)


def extract_text_from_docx(filepath: str) -> str:
    """Trích xuất toàn bộ văn bản từ file .docx bằng python-docx hoặc docx2txt."""
    try:
        import docx
        doc = docx.Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception:
        import docx2txt
        return docx2txt.process(filepath) or ""


def load_docx_files(directory: str) -> list[dict]:
    """Đọc tất cả file .docx trong thư mục, trả về list {content, source}."""
    documents = []
    docx_files = glob.glob(os.path.join(directory, "*.docx"))

    if not docx_files:
        print(f"⚠️  Không tìm thấy file .docx nào trong: {directory}")
        return documents

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        print(f"📄 Đang đọc: {filename}")
        try:
            content = extract_text_from_docx(filepath)
            if content and content.strip():
                documents.append({
                    "content": content.strip(),
                    "source": filename,
                })
                print(f"   ✅ Đọc thành công ({len(content)} ký tự)")
            else:
                print(f"   ⚠️  File rỗng, bỏ qua")
        except Exception as e:
            print(f"   ❌ Lỗi đọc file: {e}")

    return documents


def chunk_documents(documents: list[dict]) -> tuple[list, list]:
    """Chia nhỏ documents thành các chunks với metadata."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
    )

    all_chunks = []
    all_metadatas = []
    chunk_counter = 0
    ingested_at = datetime.now().isoformat()

    for doc in documents:
        chunks = text_splitter.split_text(doc["content"])
        for chunk in chunks:
            all_chunks.append(chunk)
            all_metadatas.append({
                "source": doc["source"],
                "chunk_id": chunk_counter,
                "ingested_at": ingested_at,
                "type": "knowledge_base",
            })
            chunk_counter += 1

    return all_chunks, all_metadatas


def create_vector_store(chunks: list[str], metadatas: list[dict]):
    """Tạo ChromaDB vector store từ chunks."""
    print(f"\n🔧 Loading FastEmbed ONNX model: {EMBEDDING_MODEL}")
    embeddings = FastEmbedEmbeddings(
        model_name=EMBEDDING_MODEL,
    )

    # Xóa vectordb cũ nếu tồn tại (full re-ingest)
    if os.path.exists(VECTORDB_DIR):
        import shutil
        shutil.rmtree(VECTORDB_DIR)
        print("🗑️  Đã xóa vector store cũ")

    print(f"📦 Đang tạo vector store với {len(chunks)} chunks...")

    # Tạo IDs cho mỗi chunk
    ids = [f"kb_chunk_{i}" for i in range(len(chunks))]

    vector_store = Chroma.from_texts(
        texts=chunks,
        metadatas=metadatas,
        embedding=embeddings,
        collection_name=CHROMA_COLLECTION_NAME,
        persist_directory=VECTORDB_DIR,
        ids=ids,
    )

    return vector_store


def main():
    print("=" * 60)
    print("🚀 ĐÔNG ĐÔ CS - Document Ingestion Pipeline")
    print("=" * 60)

    # Step 1: Load documents
    print(f"\n📂 Thư mục tài liệu: {DOCUMENTS_DIR}")
    documents = load_docx_files(DOCUMENTS_DIR)

    if not documents:
        print("\n❌ Không có tài liệu nào để xử lý. Kết thúc.")
        return

    print(f"\n📊 Tổng cộng: {len(documents)} file(s)")

    # Step 2: Chunking
    print("\n✂️  Đang chia nhỏ tài liệu...")
    chunks, metadatas = chunk_documents(documents)
    print(f"   → Tạo được {len(chunks)} chunks")

    # Step 3: Embedding & Store
    vector_store = create_vector_store(chunks, metadatas)

    # Step 4: Verify
    print("\n🔍 Kiểm tra vector store...")
    test_results = vector_store.similarity_search("nạp tiền", k=3)
    print(f"   → Test query 'nạp tiền': tìm được {len(test_results)} kết quả")
    for i, doc in enumerate(test_results):
        preview = doc.page_content[:100].replace("\n", " ")
        print(f"   [{i+1}] {preview}...")

    print("\n" + "=" * 60)
    print("✅ INGEST HOÀN TẤT!")
    print(f"   📦 Vector store: {VECTORDB_DIR}")
    print(f"   📊 Tổng chunks: {len(chunks)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
